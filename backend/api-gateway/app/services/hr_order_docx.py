"""Заполнение шаблона приказа о приёме (DOCX, плейсхолдеры в квадратных скобках)."""

from __future__ import annotations

import io
import json
import re
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_HIRE = ROOT / "resources" / "hr_templates" / "order_hire.docx"


def _fmt_d(d: date | None) -> str:
    if d is None:
        return "—"
    return f"{d.day:02d}.{d.month:02d}.{d.year}"


def _parse_meta(raw: str | None) -> dict[str, Any]:
    if not raw:
        return {}
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {}


def _initials_from_full_name(full_name: str) -> str:
    """И.О. Фамилия из «Фамилия Имя Отчество» (эвристика)."""
    parts = [p for p in re.split(r"\s+", full_name.strip()) if p]
    if len(parts) >= 3:
        fam, first, pat = parts[0], parts[1], parts[2]
        return f"{first[0]}.{pat[0]}. {fam}"
    if len(parts) == 2:
        return f"{parts[1][0]}. {parts[0]}"
    return full_name


def build_hire_order_mapping(
    *,
    organization_short_name: str,
    employee_full_name: str,
    position_title: str,
    hire_date: date,
    hr_meta: dict[str, Any],
    city: str,
    director_initials: str,
    employee_initials: str | None,
    application_number: str | None,
) -> dict[str, str]:
    hire_order_date_raw = hr_meta.get("hire_order_date")
    hire_order_date: date | None = None
    if isinstance(hire_order_date_raw, str) and hire_order_date_raw:
        try:
            hire_order_date = date.fromisoformat(hire_order_date_raw[:10])
        except ValueError:
            hire_order_date = None
    if hire_order_date is None:
        hire_order_date = hire_date

    hire_order_number = str(hr_meta.get("hire_order_number") or "").strip() or "—"

    cn_raw = hr_meta.get("contract_number")
    contract_number = str(cn_raw).strip() if cn_raw else ""
    cd_raw = hr_meta.get("contract_date")
    contract_date: date | None = None
    if isinstance(cd_raw, str) and cd_raw:
        try:
            contract_date = date.fromisoformat(cd_raw[:10])
        except ValueError:
            contract_date = None

    if contract_number and contract_date:
        contract_phrase = f"трудового договора № {contract_number} от {_fmt_d(contract_date)}"
    elif contract_number:
        contract_phrase = f"трудового договора № {contract_number}"
    else:
        contract_phrase = "трудового договора"

    app_no = (application_number or "").strip() or "б/н"

    emp_ini = (employee_initials or "").strip() or _initials_from_full_name(employee_full_name)

    return {
        "[Краткое название организации]": organization_short_name.strip() or "—",
        "[Дата приказа]": _fmt_d(hire_order_date),
        "[Номер приказа]": hire_order_number,
        "[Город]": city.strip() or "—",
        "[ФИО сотрудника]": employee_full_name.strip(),
        "[Должность]": position_title.strip() or "—",
        "[Дата приема]": _fmt_d(hire_date),
        "[трудовой договор/контракт]": contract_phrase,
        "[Номер заявления]": app_no,
        "[Дата контракта]": _fmt_d(contract_date) if contract_date else _fmt_d(hire_order_date),
        "[Номер контракта]": contract_number or "—",
        "[Инициалы, фамилия директора]": director_initials.strip() or "—",
        "[Инициалы, фамилия сотрудника]": emp_ini,
    }


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    t = paragraph.text
    if not t or "[" not in t:
        return
    orig = t
    for old, new in mapping.items():
        if old in t:
            t = t.replace(old, new)
    if t != orig:
        paragraph.text = t


def _walk_cell(cell, mapping: dict[str, str]) -> None:
    for p in cell.paragraphs:
        _replace_in_paragraph(p, mapping)
    for tbl in cell.tables:
        _walk_table(tbl, mapping)


def _walk_table(table, mapping: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            _walk_cell(cell, mapping)


def render_hire_order_docx(mapping: dict[str, str]) -> bytes:
    if not TEMPLATE_HIRE.is_file():
        raise FileNotFoundError(f"Шаблон не найден: {TEMPLATE_HIRE}")
    doc = Document(str(TEMPLATE_HIRE))
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for tbl in doc.tables:
        _walk_table(tbl, mapping)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
